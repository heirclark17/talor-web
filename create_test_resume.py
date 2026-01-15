#!/usr/bin/env python3
"""Create a test resume for end-to-end testing"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Header - Name
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name.add_run('SARAH CHEN')
name_run.font.size = Pt(18)
name_run.font.bold = True
name_run.font.color.rgb = RGBColor(0, 51, 102)

# Contact Info
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact.add_run('Seattle, WA | (206) 555-0123 | sarah.chen@email.com | linkedin.com/in/sarahchen')
contact_run.font.size = Pt(10)

# Professional Summary
summary_heading = doc.add_paragraph()
summary_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
summary_run.font.size = Pt(12)
summary_run.font.bold = True
summary_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph(
    'Security Delivery Practice Manager with 12+ years driving security programs, risk management, and '
    'cross-functional team leadership in cloud-native and enterprise environments. Expertise in building '
    'and scaling security delivery frameworks, implementing DevSecOps practices, and managing complex '
    'security initiatives across global teams. Proven track record of reducing security incidents by 45 percent '
    'while accelerating delivery velocity by 30 percent through automation and process optimization.'
)

# Core Competencies
comp_heading = doc.add_paragraph()
comp_run = comp_heading.add_run('CORE COMPETENCIES')
comp_run.font.size = Pt(12)
comp_run.font.bold = True
comp_run.font.color.rgb = RGBColor(0, 51, 102)

# Create competencies table
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'

competencies = [
    'Security Program Leadership', 'AWS Security Architecture', 'DevSecOps and CI/CD Security',
    'Risk Management', 'Security Delivery Frameworks', 'Cross-Functional Leadership',
    'Threat Modeling', 'Compliance (SOC 2, ISO 27001)', 'Security Champions Program',
    'Incident Response', 'Vulnerability Management', 'Agile and Scrum'
]

for i, comp in enumerate(competencies):
    row = i // 3
    col = i % 3
    cell = table.rows[row].cells[col]
    cell.text = comp

# Professional Experience
exp_heading = doc.add_paragraph()
exp_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
exp_run.font.size = Pt(12)
exp_run.font.bold = True
exp_run.font.color.rgb = RGBColor(0, 51, 102)

# Job 1
job1 = doc.add_paragraph()
job1_run = job1.add_run('Senior Security Program Manager')
job1_run.font.size = Pt(11)
job1_run.font.bold = True

company1 = doc.add_paragraph()
company1_run = company1.add_run('Microsoft Corporation | Seattle, WA | June 2020 - Present')
company1_run.font.size = Pt(10)
company1_run.font.italic = True

doc.add_paragraph('Lead security delivery program for Azure DevOps platform serving 85,000 plus enterprise customers, managing cross-functional teams of 40 plus engineers, security analysts, and compliance specialists across 5 time zones')
doc.add_paragraph('Architected and implemented DevSecOps framework integrating automated security scanning into CI/CD pipelines, reducing critical vulnerabilities in production by 67 percent and achieving 99.8 percent deployment success rate')
doc.add_paragraph('Built and scaled Security Champions program across 12 product teams, training 150 plus developers on secure coding practices, threat modeling, and OWASP Top 10, resulting in 52 percent reduction in security defects')
doc.add_paragraph('Managed FedRAMP High authorization program for government cloud platform, coordinating vulnerability remediation using NIST 800-53 controls and achieving Authority to Operate with zero critical findings')
doc.add_paragraph('Drove security incident response for 200 plus security events annually, reducing mean time to resolution from 8 hours to 2.5 hours through automation and playbook optimization')

# Job 2
doc.add_paragraph()
job2 = doc.add_paragraph()
job2_run = job2.add_run('Security Program Manager')
job2_run.font.size = Pt(11)
job2_run.font.bold = True

company2 = doc.add_paragraph()
company2_run = company2.add_run('Salesforce | San Francisco, CA | January 2018 - May 2020')
company2_run.font.size = Pt(10)
company2_run.font.italic = True

doc.add_paragraph('Managed cloud security program for Salesforce Commerce Cloud, overseeing vulnerability management, penetration testing, and security assessments for 100 plus microservices')
doc.add_paragraph('Implemented threat modeling framework across 25 product teams, conducting 80 plus threat model reviews annually and identifying 300 plus security requirements that prevented high-impact vulnerabilities')
doc.add_paragraph('Led SOC 2 Type II and ISO 27001 certification programs, coordinating audit activities, remediating 95 plus control deficiencies, and achieving clean audit opinions with zero findings')

# Job 3
doc.add_paragraph()
job3 = doc.add_paragraph()
job3_run = job3.add_run('Senior Security Engineer')
job3_run.font.size = Pt(11)
job3_run.font.bold = True

company3 = doc.add_paragraph()
company3_run = company3.add_run('Cisco Systems | San Jose, CA | March 2013 - December 2017')
company3_run.font.size = Pt(10)
company3_run.font.italic = True

doc.add_paragraph('Led security architecture and implementation for enterprise SaaS platform with 500,000 plus users, designing secure authentication, authorization, and data protection mechanisms')
doc.add_paragraph('Managed vulnerability management program coordinating security scanning, penetration testing, and remediation tracking for 1,500 plus assets across cloud and on-premise infrastructure')

# Education
doc.add_paragraph()
edu_heading = doc.add_paragraph()
edu_run = edu_heading.add_run('EDUCATION')
edu_run.font.size = Pt(12)
edu_run.font.bold = True
edu_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('Master of Science in Cybersecurity | University of Washington | 2013')
doc.add_paragraph('Bachelor of Science in Computer Science | UC Berkeley | 2011')

# Certifications
cert_heading = doc.add_paragraph()
cert_run = cert_heading.add_run('CERTIFICATIONS')
cert_run.font.size = Pt(12)
cert_run.font.bold = True
cert_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('CISSP - Certified Information Systems Security Professional (2016)')
doc.add_paragraph('AWS Certified Security - Specialty (2021)')
doc.add_paragraph('Certified Cloud Security Professional (CCSP) - 2019')
doc.add_paragraph('Certified Scrum Master (CSM) - 2018')

# Save document
doc.save('Sarah_Chen_Security_Resume.docx')
print('Resume created successfully: Sarah_Chen_Security_Resume.docx')

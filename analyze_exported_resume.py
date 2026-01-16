"""
Analyze exported DOCX resume for ATS compliance and missing fields
Based on 2026 ATS best practices research
"""

import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def analyze_docx(file_path):
    """Analyze DOCX file for ATS compliance"""

    print("=" * 80)
    print("ATS COMPLIANCE ANALYSIS - EXPORTED RESUME")
    print("=" * 80)
    print(f"\nAnalyzing: {file_path}\n")

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"ERROR: Could not open file: {e}")
        return

    issues = []
    warnings = []
    passed = []

    # Extract all text content
    all_text = '\n'.join([para.text for para in doc.paragraphs])

    print("\n" + "="*80)
    print("1. DOCUMENT PROPERTIES")
    print("="*80)

    # Check document properties
    props = doc.core_properties
    print(f"Title: {props.title or '❌ MISSING'}")
    print(f"Author: {props.author or '❌ MISSING'}")
    print(f"Subject: {props.subject or '❌ MISSING'}")

    if not props.title:
        issues.append("Missing document title property")
    if not props.author:
        issues.append("Missing document author property")

    print("\n" + "="*80)
    print("2. PAGE SETUP")
    print("="*80)

    # Check margins
    for i, section in enumerate(doc.sections):
        print(f"\nSection {i+1}:")
        top_margin = section.top_margin.inches if section.top_margin else 0
        bottom_margin = section.bottom_margin.inches if section.bottom_margin else 0
        left_margin = section.left_margin.inches if section.left_margin else 0
        right_margin = section.right_margin.inches if section.right_margin else 0

        print(f"  Top margin: {top_margin:.2f} inches")
        print(f"  Bottom margin: {bottom_margin:.2f} inches")
        print(f"  Left margin: {left_margin:.2f} inches")
        print(f"  Right margin: {right_margin:.2f} inches")

        # ATS best practice: 0.75-1 inch margins
        if top_margin < 0.5 or top_margin > 1.5:
            warnings.append(f"Top margin ({top_margin:.2f}\") outside recommended range (0.75-1.0\")")
        if left_margin < 0.5 or left_margin > 1.5:
            warnings.append(f"Left margin ({left_margin:.2f}\") outside recommended range (0.75-1.0\")")

        # Check for headers/footers
        if section.header.paragraphs and any(p.text.strip() for p in section.header.paragraphs):
            issues.append("❌ CRITICAL: Content found in header (25% of ATS systems cannot read headers)")
        else:
            passed.append("✓ No content in headers")

        if section.footer.paragraphs and any(p.text.strip() for p in section.footer.paragraphs):
            issues.append("❌ CRITICAL: Content found in footer (25% of ATS systems cannot read footers)")
        else:
            passed.append("✓ No content in footers")

    print("\n" + "="*80)
    print("3. REQUIRED CONTACT INFORMATION")
    print("="*80)

    # Check for contact information (first 10 lines)
    first_lines = '\n'.join([para.text for para in doc.paragraphs[:10]])

    # Check for name (should be in first few lines, larger font)
    name_found = False
    for para in doc.paragraphs[:5]:
        if para.text.strip() and len(para.text.strip()) > 5:
            if para.runs and para.runs[0].font.size:
                size = para.runs[0].font.size.pt if para.runs[0].font.size else 12
                if size >= 16:  # Name should be 18-24pt
                    print(f"✓ Name found: '{para.text}' ({size}pt)")
                    name_found = True
                    passed.append("✓ Name present with large font size")
                    break

    if not name_found:
        issues.append("❌ CRITICAL: Candidate name not found or not properly formatted (should be 18-24pt)")

    # Check for phone number
    phone_pattern = r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    if re.search(phone_pattern, first_lines):
        phone_match = re.search(phone_pattern, first_lines)
        print(f"✓ Phone number found: {phone_match.group()}")
        passed.append("✓ Phone number present")
    else:
        issues.append("❌ CRITICAL: Phone number not found in contact section")

    # Check for email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    if re.search(email_pattern, first_lines):
        email_match = re.search(email_pattern, first_lines)
        print(f"✓ Email found: {email_match.group()}")
        passed.append("✓ Email address present")
    else:
        issues.append("❌ CRITICAL: Email address not found in contact section")

    # Check for LinkedIn
    linkedin_patterns = [r'linkedin\.com/in/[\w-]+', r'linkedin\.com', r'LinkedIn']
    linkedin_found = any(re.search(pattern, first_lines, re.IGNORECASE) for pattern in linkedin_patterns)
    if linkedin_found:
        print("✓ LinkedIn profile mentioned")
        passed.append("✓ LinkedIn profile present")
    else:
        warnings.append("LinkedIn profile not found (recommended but optional)")

    # Check for location
    location_pattern = r'[A-Z][a-z]+,\s*[A-Z]{2}'  # City, ST format
    if re.search(location_pattern, first_lines):
        location_match = re.search(location_pattern, first_lines)
        print(f"✓ Location found: {location_match.group()}")
        passed.append("✓ Location present")
    else:
        warnings.append("Location not found in standard format (City, ST)")

    print("\n" + "="*80)
    print("4. REQUIRED RESUME SECTIONS")
    print("="*80)

    # Check for standard section headings
    required_sections = {
        'Professional Summary': [r'PROFESSIONAL SUMMARY', r'SUMMARY', r'PROFILE', r'CAREER SUMMARY'],
        'Experience': [r'PROFESSIONAL EXPERIENCE', r'WORK EXPERIENCE', r'EXPERIENCE', r'EMPLOYMENT HISTORY'],
        'Skills': [r'CORE COMPETENCIES', r'SKILLS', r'TECHNICAL SKILLS', r'KEY SKILLS'],
        'Education': [r'EDUCATION'],
    }

    for section_name, patterns in required_sections.items():
        found = False
        for pattern in patterns:
            if re.search(pattern, all_text, re.IGNORECASE):
                print(f"✓ {section_name} section found")
                passed.append(f"✓ {section_name} section present")
                found = True
                break

        if not found:
            issues.append(f"❌ CRITICAL: {section_name} section not found (ATS requires standard headings)")

    # Optional sections
    optional_sections = {
        'Certifications': [r'CERTIFICATIONS', r'CERTIFICATION', r'TRAINING', r'PROFESSIONAL DEVELOPMENT'],
    }

    for section_name, patterns in optional_sections.items():
        found = any(re.search(pattern, all_text, re.IGNORECASE) for pattern in patterns)
        if found:
            print(f"✓ {section_name} section found (optional)")
            passed.append(f"✓ {section_name} section present")
        else:
            print(f"  {section_name} section not found (optional)")

    print("\n" + "="*80)
    print("5. EXPERIENCE SECTION ANALYSIS")
    print("="*80)

    # Check for job titles and dates in experience section
    experience_section_start = None
    experience_section_end = None

    for i, para in enumerate(doc.paragraphs):
        text_upper = para.text.upper()
        if any(pattern in text_upper for pattern in ['PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE', 'EXPERIENCE']):
            experience_section_start = i
        elif experience_section_start and any(pattern in text_upper for pattern in ['EDUCATION', 'CERTIFICATIONS', 'SKILLS']):
            experience_section_end = i
            break

    if experience_section_start:
        exp_paragraphs = doc.paragraphs[experience_section_start+1:experience_section_end] if experience_section_end else doc.paragraphs[experience_section_start+1:]
        exp_text = '\n'.join([p.text for p in exp_paragraphs])

        # Check for dates (MM/YYYY or Month YYYY format)
        date_patterns = [
            r'\d{2}/\d{4}',  # MM/YYYY
            r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{4}',  # Month YYYY
            r'Present'
        ]

        dates_found = []
        for pattern in date_patterns:
            matches = re.findall(pattern, exp_text)
            dates_found.extend(matches)

        print(f"Dates found in experience section: {len(dates_found)}")
        if dates_found:
            print(f"  Sample dates: {dates_found[:5]}")
            passed.append(f"✓ Employment dates present ({len(dates_found)} found)")
        else:
            issues.append("❌ CRITICAL: No employment dates found in experience section")

        # Check for job titles (look for bold text or specific patterns)
        job_title_count = 0
        for para in exp_paragraphs:
            if para.runs and para.runs[0].font.bold and len(para.text) > 5:
                job_title_count += 1
                if job_title_count <= 3:
                    print(f"  Job title found: '{para.text.strip()}'")

        if job_title_count > 0:
            print(f"✓ {job_title_count} job titles found (bold formatting)")
            passed.append(f"✓ {job_title_count} job titles found")
        else:
            issues.append("❌ CRITICAL: No job titles found (should be bold)")

        # Check for company names and locations
        company_location_pattern = r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,?\s+[A-Z][a-z]+(?:,\s*[A-Z]{2})?'
        companies = re.findall(company_location_pattern, exp_text)
        if companies:
            print(f"✓ Potential company/location entries: {len(companies)}")
            print(f"  Sample: {companies[:3]}")

        # Check for bullet points
        bullet_count = sum(1 for para in exp_paragraphs if para.text.strip().startswith('•') or para.text.strip().startswith('-'))
        print(f"Bullet points in experience: {bullet_count}")
        if bullet_count > 0:
            passed.append(f"✓ {bullet_count} bullet points in experience")
        else:
            warnings.append("No bullet points found in experience (recommended format)")

    print("\n" + "="*80)
    print("6. FONT ANALYSIS")
    print("="*80)

    fonts_used = set()
    font_sizes = []

    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts_used.add(run.font.name)
            if run.font.size:
                font_sizes.append(run.font.size.pt)

    print("Fonts used in document:")
    for font in sorted(fonts_used):
        print(f"  - {font}")

    # ATS-approved fonts
    ats_approved = ['Calibri', 'Arial', 'Times New Roman', 'Helvetica', 'Georgia', 'Verdana', 'Garamond']
    non_approved = [f for f in fonts_used if f not in ats_approved]

    if non_approved:
        warnings.append(f"Non-standard fonts used: {', '.join(non_approved)} (may not be ATS-friendly)")
    else:
        passed.append("✓ All fonts are ATS-approved")

    if font_sizes:
        print(f"\nFont size range: {min(font_sizes):.1f}pt - {max(font_sizes):.1f}pt")
        if min(font_sizes) < 10:
            issues.append(f"❌ Font size too small: {min(font_sizes):.1f}pt (minimum 10pt required for ATS)")
        else:
            passed.append("✓ All font sizes >= 10pt")

    print("\n" + "="*80)
    print("7. LAYOUT ANALYSIS")
    print("="*80)

    # Check for tables
    table_count = len(doc.tables)
    print(f"Tables in document: {table_count}")
    if table_count > 0:
        issues.append(f"❌ {table_count} table(s) found - tables can cause ATS parsing errors")
        for i, table in enumerate(doc.tables):
            print(f"  Table {i+1}: {len(table.rows)} rows x {len(table.columns)} columns")
    else:
        passed.append("✓ No tables (single-column layout recommended)")

    # Check for images
    # (Note: Images embedded via runs need special handling)
    print("Image detection: (basic check)")
    print("  Note: Full image detection requires additional analysis")

    # Check alignment
    alignments = []
    for para in doc.paragraphs:
        if para.alignment:
            alignments.append(para.alignment)

    left_aligned = sum(1 for a in alignments if a == WD_ALIGN_PARAGRAPH.LEFT)
    if left_aligned > len(alignments) * 0.8:  # 80% left-aligned
        passed.append("✓ Most content is left-aligned (ATS-friendly)")
    else:
        warnings.append("Some content may not be left-aligned (ATS prefers left alignment)")

    print("\n" + "="*80)
    print("ANALYSIS SUMMARY")
    print("="*80)

    print(f"\n✓ PASSED CHECKS: {len(passed)}")
    for item in passed:
        print(f"  {item}")

    print(f"\n⚠ WARNINGS: {len(warnings)}")
    for item in warnings:
        print(f"  {item}")

    print(f"\n❌ CRITICAL ISSUES: {len(issues)}")
    for item in issues:
        print(f"  {item}")

    # Calculate ATS compliance score
    total_checks = len(passed) + len(warnings) + len(issues)
    compliance_score = (len(passed) / total_checks * 100) if total_checks > 0 else 0

    print(f"\n" + "="*80)
    print(f"ATS COMPLIANCE SCORE: {compliance_score:.1f}%")

    if compliance_score >= 90:
        print("Rating: EXCELLENT - Ready for ATS submission")
    elif compliance_score >= 75:
        print("Rating: GOOD - Minor improvements recommended")
    elif compliance_score >= 60:
        print("Rating: FAIR - Several improvements needed")
    else:
        print("Rating: POOR - Major revisions required for ATS compatibility")

    print("="*80)

    # Return summary
    return {
        'compliance_score': compliance_score,
        'passed': len(passed),
        'warnings': len(warnings),
        'critical_issues': len(issues),
        'issues_list': issues
    }

if __name__ == '__main__':
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        # Default to exported file
        file_path = r'C:\Users\derri\Downloads\exported-tailored-resume.docx'

    analyze_docx(file_path)

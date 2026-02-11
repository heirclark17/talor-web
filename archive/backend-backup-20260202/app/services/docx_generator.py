from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re
from datetime import datetime
from app.config import get_settings

settings = get_settings()

class DOCXGenerator:
    """Generate formatted DOCX resumes"""

    def __init__(self):
        self.company_colors = {
            "jpmorgan": RGBColor(0, 51, 102),  # Navy blue
            "oracle": RGBColor(255, 0, 0),  # Oracle red
            "microsoft": RGBColor(0, 103, 184),  # Microsoft blue
            "amazon": RGBColor(255, 153, 0),  # Amazon orange
            "google": RGBColor(66, 133, 244),  # Google blue
            "default": RGBColor(0, 51, 102)  # Navy blue default
        }

    def get_company_color(self, company_name: str) -> RGBColor:
        """Get company brand color"""
        company_lower = company_name.lower()
        for key, color in self.company_colors.items():
            if key in company_lower:
                return color
        return self.company_colors["default"]

    def add_hyperlink(self, paragraph, url: str, text: str):
        """Add clickable hyperlink to paragraph"""
        # This creates the  w:hyperlink tag and add needed values
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create a new run with the text
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Set color to blue and underline
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)

        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)
        new_run.text = text

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

        return hyperlink

    def create_tailored_resume(
        self,
        candidate_name: str,
        contact_info: dict,
        job_details: dict,
        tailored_content: dict,
        base_resume_data: dict,
        output_filename: str
    ) -> str:
        """
        Create a tailored resume DOCX file

        Args:
            candidate_name: Full name
            contact_info: {email, phone, location, linkedin}
            job_details: {company, title, url}
            tailored_content: {summary, experience, competencies, alignment_statement}
            base_resume_data: {education, certifications}
            output_filename: Filename for output

        Returns:
            Path to generated DOCX file
        """

        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        company_color = self.get_company_color(job_details.get('company', ''))

        # Target Position Box
        box = doc.add_paragraph()
        box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        box_run = box.add_run("TARGET POSITION\n")
        box_run.bold = True
        box_run.font.size = Pt(11)

        title_run = box.add_run(f"{job_details.get('title', 'Position')}\n")
        title_run.bold = True
        title_run.font.size = Pt(12)

        company_run = box.add_run(f"{job_details.get('company', 'Company')} | {contact_info.get('location', 'Location')}\n")
        company_run.font.size = Pt(11)

        # Add application link if URL provided
        if job_details.get('url'):
            link_para = doc.add_paragraph()
            link_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.add_hyperlink(link_para, job_details['url'], "Apply Here")

        doc.add_paragraph()  # Spacing

        # Header - Name and Contact
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = header.add_run(candidate_name + "\n")
        name_run.bold = True
        name_run.font.size = Pt(18)

        contact_run = header.add_run(
            f"{contact_info.get('email', '')} | {contact_info.get('phone', '')} | {contact_info.get('location', '')}"
        )
        contact_run.font.size = Pt(10)

        doc.add_paragraph()  # Spacing

        # Professional Summary
        self._add_section_heading(doc, "PROFESSIONAL SUMMARY", company_color)
        summary_para = doc.add_paragraph(tailored_content.get('summary', ''))
        summary_para.paragraph_format.line_spacing = 1.15

        # Core Competencies
        self._add_section_heading(doc, "CORE COMPETENCIES", company_color)
        competencies = tailored_content.get('competencies', [])

        if competencies:
            # Create 3-column table
            table = doc.add_table(rows=(len(competencies) + 2) // 3, cols=3)
            table.style = 'Table Grid'

            for idx, comp in enumerate(competencies):
                row = idx // 3
                col = idx % 3
                cell = table.cell(row, col)
                cell.text = f"• {comp}"
                cell.paragraphs[0].runs[0].font.size = Pt(10)

        # Professional Experience
        self._add_section_heading(doc, "PROFESSIONAL EXPERIENCE", company_color)

        for exp in tailored_content.get('experience', []):
            # Job header
            job_header = doc.add_paragraph()
            header_run = job_header.add_run(exp.get('header', ''))
            header_run.bold = True
            header_run.font.size = Pt(11)

            # Bullets
            for bullet in exp.get('bullets', []):
                # Skip empty bullets and bullets containing only separators/whitespace
                # This regex catches any combination of whitespace and separator characters
                if bullet and bullet.strip() and not re.match(r'^[\s\|\/•\-–—]+$', bullet.strip()):
                    bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                    bullet_para.paragraph_format.line_spacing = 1.15
                    bullet_para.paragraph_format.space_after = Pt(6)

            doc.add_paragraph()  # Spacing between jobs

        # Education
        if base_resume_data.get('education'):
            self._add_section_heading(doc, "EDUCATION", company_color)
            edu_para = doc.add_paragraph(base_resume_data['education'])
            edu_para.paragraph_format.line_spacing = 1.15

        # Certifications
        if base_resume_data.get('certifications'):
            self._add_section_heading(doc, "CERTIFICATIONS & TRAINING", company_color)
            cert_para = doc.add_paragraph(base_resume_data['certifications'])
            cert_para.paragraph_format.line_spacing = 1.15

        # Alignment Statement
        if tailored_content.get('alignment_statement'):
            self._add_section_heading(doc, f"ALIGNMENT WITH {job_details.get('company', 'COMPANY').upper()} MISSION", company_color)
            align_para = doc.add_paragraph(tailored_content['alignment_statement'])
            align_para.paragraph_format.line_spacing = 1.15
            for run in align_para.runs:
                run.font.size = Pt(10)

        # Save document
        os.makedirs(settings.resumes_dir, exist_ok=True)
        output_path = os.path.join(settings.resumes_dir, output_filename)
        doc.save(output_path)

        return output_path

    def _add_section_heading(self, doc, text: str, color: RGBColor):
        """Add a section heading with color"""
        heading = doc.add_paragraph()
        heading_run = heading.add_run(text)
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        heading_run.font.color.rgb = color
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

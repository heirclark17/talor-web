"""
Resume Export Service - Generate PDF and DOCX files

Creates properly formatted resume documents with correct filenames:
UserName_TargetRole_TailoredResume.pdf
UserName_TargetRole_TailoredResume.docx
"""

import io
import re
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors

class ResumeExportService:
    def __init__(self):
        pass

    def generate_filename(self, user_name: str, target_role: str, extension: str) -> str:
        """
        Generate proper filename for resume export
        Format: UserName_TargetRole_TailoredResume.ext
        """
        # Clean name - remove special characters, replace spaces with underscores
        clean_name = re.sub(r'[^\w\s-]', '', user_name).strip()
        clean_name = re.sub(r'[-\s]+', '_', clean_name)

        # Clean role - remove special characters, replace spaces with underscores
        clean_role = re.sub(r'[^\w\s-]', '', target_role).strip()
        clean_role = re.sub(r'[-\s]+', '_', clean_role)

        # Format: UserName_TargetRole_TailoredResume.ext
        filename = f"{clean_name}_{clean_role}_TailoredResume.{extension}"

        return filename

    def generate_docx(
        self,
        resume_data: Dict[str, Any],
        user_name: str,
        target_role: str
    ) -> io.BytesIO:
        """
        Generate DOCX file from resume data
        """
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Add name (header)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(user_name)
        name_run.font.size = Pt(18)
        name_run.font.bold = True
        name_run.font.name = 'Calibri'

        # Add contact info
        contact_info = resume_data.get('contact', {})
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_text = f"{contact_info.get('email', '')} | {contact_info.get('phone', '')} | {contact_info.get('location', '')}"
        contact_run = contact_para.add_run(contact_text)
        contact_run.font.size = Pt(10)
        contact_run.font.name = 'Calibri'

        doc.add_paragraph()  # Spacer

        # Add Professional Summary
        if resume_data.get('summary'):
            summary_heading = doc.add_paragraph()
            summary_heading_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
            summary_heading_run.font.size = Pt(12)
            summary_heading_run.font.bold = True
            summary_heading_run.font.color.rgb = RGBColor(0, 51, 102)

            summary_para = doc.add_paragraph(resume_data['summary'])
            summary_para.style = 'Normal'

            doc.add_paragraph()  # Spacer

        # Add Experience
        if resume_data.get('experience'):
            exp_heading = doc.add_paragraph()
            exp_heading_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
            exp_heading_run.font.size = Pt(12)
            exp_heading_run.font.bold = True
            exp_heading_run.font.color.rgb = RGBColor(0, 51, 102)

            experiences = resume_data['experience']
            if isinstance(experiences, list):
                for exp in experiences:
                    # Handle both formats: {header} or {title, company}
                    if 'header' in exp:
                        # Format from parser: "Job Title – Company Name"
                        job_title_para = doc.add_paragraph()
                        job_title_run = job_title_para.add_run(exp.get('header', ''))
                        job_title_run.font.size = Pt(11)
                        job_title_run.font.bold = True
                    else:
                        # Format from tailor service: separate title and company
                        job_title_para = doc.add_paragraph()
                        job_title_run = job_title_para.add_run(exp.get('title', ''))
                        job_title_run.font.size = Pt(11)
                        job_title_run.font.bold = True

                    # Add company/location/dates line (only if not empty)
                    if 'header' in exp:
                        # Just location and dates for header format
                        company_text = f"{exp.get('location', '')} | {exp.get('dates', '')}"
                    else:
                        # Full company, location, dates for split format
                        company_text = f"{exp.get('company', '')} | {exp.get('location', '')} | {exp.get('dates', '')}"

                    # Only add paragraph if company_text is not just separators/whitespace
                    if company_text and not re.match(r'^[\s\|\/•\-–—]+$', company_text):
                        company_para = doc.add_paragraph()
                        company_run = company_para.add_run(company_text)
                        company_run.font.italic = True

                    # Add bullet points
                    if exp.get('bullets'):
                        for bullet in exp['bullets']:
                            # Skip empty bullets and bullets containing only separators/whitespace
                            # This regex catches any combination of whitespace and separator characters
                            if bullet and bullet.strip() and not re.match(r'^[\s\|\/•\-–—]+$', bullet.strip()):
                                doc.add_paragraph(bullet, style='List Bullet')

                    doc.add_paragraph()  # Spacer between jobs

        # Add Skills
        if resume_data.get('skills'):
            skills_heading = doc.add_paragraph()
            skills_heading_run = skills_heading.add_run('CORE COMPETENCIES')
            skills_heading_run.font.size = Pt(12)
            skills_heading_run.font.bold = True
            skills_heading_run.font.color.rgb = RGBColor(0, 51, 102)

            skills = resume_data['skills']
            if isinstance(skills, list):
                skills_text = ' • '.join(skills)
                doc.add_paragraph(skills_text)
            elif isinstance(skills, str):
                doc.add_paragraph(skills)

            doc.add_paragraph()  # Spacer

        # Add Education
        if resume_data.get('education'):
            edu_heading = doc.add_paragraph()
            edu_heading_run = edu_heading.add_run('EDUCATION')
            edu_heading_run.font.size = Pt(12)
            edu_heading_run.font.bold = True
            edu_heading_run.font.color.rgb = RGBColor(0, 51, 102)

            education = resume_data['education']
            if isinstance(education, list):
                for edu in education:
                    edu_para = doc.add_paragraph()
                    edu_text = f"{edu.get('degree', '')} | {edu.get('institution', '')} | {edu.get('year', '')}"
                    edu_para.add_run(edu_text)
            elif isinstance(education, str):
                doc.add_paragraph(education)

        # Save to BytesIO
        file_buffer = io.BytesIO()
        doc.save(file_buffer)
        file_buffer.seek(0)

        return file_buffer

    def generate_pdf(
        self,
        resume_data: Dict[str, Any],
        user_name: str,
        target_role: str
    ) -> io.BytesIO:
        """
        Generate PDF file from resume data
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch
        )

        # Container for the 'Flowable' objects
        elements = []

        # Define styles
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#000000'),
            spaceAfter=6,
            alignment=1  # Center
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor=colors.HexColor('#003366'),
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )

        body_style = styles['BodyText']
        body_style.fontSize = 10
        body_style.spaceAfter = 6

        # Add name
        elements.append(Paragraph(user_name, title_style))

        # Add contact info
        contact_info = resume_data.get('contact', {})
        contact_text = f"{contact_info.get('email', '')} | {contact_info.get('phone', '')} | {contact_info.get('location', '')}"
        elements.append(Paragraph(contact_text, body_style))
        elements.append(Spacer(1, 12))

        # Add Professional Summary
        if resume_data.get('summary'):
            elements.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
            elements.append(Paragraph(resume_data['summary'], body_style))
            elements.append(Spacer(1, 12))

        # Add Experience
        if resume_data.get('experience'):
            elements.append(Paragraph('PROFESSIONAL EXPERIENCE', heading_style))

            experiences = resume_data['experience']
            if isinstance(experiences, list):
                for exp in experiences:
                    # Handle both formats: {header} or {title, company}
                    if 'header' in exp:
                        # Format from parser: "Job Title – Company Name"
                        job_title = f"<b>{exp.get('header', '')}</b>"
                        elements.append(Paragraph(job_title, body_style))

                        # Just location and dates for header format
                        company_text = f"{exp.get('location', '')} | {exp.get('dates', '')}"
                        elements.append(Paragraph(f"<i>{company_text}</i>", body_style))
                    else:
                        # Format from tailor service: separate title and company
                        job_title = f"<b>{exp.get('title', '')}</b>"
                        elements.append(Paragraph(job_title, body_style))

                        # Full company, location, dates for split format
                        company_text = f"{exp.get('company', '')} | {exp.get('location', '')} | {exp.get('dates', '')}"
                        elements.append(Paragraph(f"<i>{company_text}</i>", body_style))

                    # Bullets
                    if exp.get('bullets'):
                        for bullet in exp['bullets']:
                            # Skip empty bullets and bullets containing only separators/whitespace
                            # This regex catches any combination of whitespace and separator characters
                            if bullet and bullet.strip() and not re.match(r'^[\s\|\/•\-–—]+$', bullet.strip()):
                                bullet_text = f"• {bullet}"
                                elements.append(Paragraph(bullet_text, body_style))

                    elements.append(Spacer(1, 6))

        # Add Skills
        if resume_data.get('skills'):
            elements.append(Paragraph('CORE COMPETENCIES', heading_style))

            skills = resume_data['skills']
            if isinstance(skills, list):
                skills_text = ' • '.join(skills)
                elements.append(Paragraph(skills_text, body_style))
            elif isinstance(skills, str):
                elements.append(Paragraph(skills, body_style))

            elements.append(Spacer(1, 12))

        # Add Education
        if resume_data.get('education'):
            elements.append(Paragraph('EDUCATION', heading_style))

            education = resume_data['education']
            if isinstance(education, list):
                for edu in education:
                    edu_text = f"{edu.get('degree', '')} | {edu.get('institution', '')} | {edu.get('year', '')}"
                    elements.append(Paragraph(edu_text, body_style))
            elif isinstance(education, str):
                elements.append(Paragraph(education, body_style))

        # Build PDF
        doc.build(elements)
        buffer.seek(0)

        return buffer
